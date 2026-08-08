"""Tests for importing plain text from supported document formats."""

from __future__ import annotations

from pathlib import Path

import pytest

from vntts.services.document_import import DocumentTextImporter
from vntts.utils.exceptions import DocumentImportError


@pytest.fixture
def importer() -> DocumentTextImporter:
    return DocumentTextImporter()


def test_imports_utf8_txt(importer: DocumentTextImporter, tmp_path: Path) -> None:
    source = tmp_path / "noi-dung.txt"
    source.write_text("Xin chào Việt Nam.\n\nĐây là đoạn thứ hai.", encoding="utf-8-sig")

    result = importer.import_file(source)

    assert result.display_name == "noi-dung.txt"
    assert result.text == "Xin chào Việt Nam.\n\nĐây là đoạn thứ hai."


def test_imports_srt_without_indices_or_timestamps(
    importer: DocumentTextImporter, tmp_path: Path
) -> None:
    source = tmp_path / "phu-de.srt"
    source.write_text(
        "1\n00:00:01,000 --> 00:00:03,000\n<i>Xin chào</i> &amp; hẹn gặp lại.\n\n"
        "2\n00:00:04,000 --> 00:00:06,000\nDòng phụ đề thứ hai.\n",
        encoding="utf-8",
    )

    result = importer.import_file(source)

    assert result.text == "Xin chào & hẹn gặp lại.\nDòng phụ đề thứ hai."
    assert "-->" not in result.text


def test_imports_docx_paragraphs_and_tables(
    importer: DocumentTextImporter, tmp_path: Path
) -> None:
    from docx import Document

    source = tmp_path / "tai-lieu.docx"
    document = Document()
    document.add_paragraph("Tiêu đề tài liệu")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cột một"
    table.cell(0, 1).text = "Cột hai"
    document.add_paragraph("Đoạn kết.")
    document.save(source)

    result = importer.import_file(source)

    assert result.text == "Tiêu đề tài liệu\nCột một\tCột hai\nĐoạn kết."


def test_imports_text_based_pdf(importer: DocumentTextImporter, tmp_path: Path) -> None:
    source = tmp_path / "tai-lieu.pdf"
    _write_minimal_pdf(source, "Xin chao tu PDF")

    result = importer.import_file(source)

    assert "Xin chao tu PDF" in result.text


def test_rejects_empty_or_image_only_pdf(
    importer: DocumentTextImporter, tmp_path: Path
) -> None:
    from pypdf import PdfWriter

    source = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with source.open("wb") as stream:
        writer.write(stream)

    with pytest.raises(DocumentImportError, match="OCR"):
        importer.import_file(source)


def test_rejects_unsupported_extension(
    importer: DocumentTextImporter, tmp_path: Path
) -> None:
    source = tmp_path / "audio.wav"
    source.write_bytes(b"RIFF")

    with pytest.raises(DocumentImportError, match="chưa được hỗ trợ"):
        importer.import_file(source)


def _write_minimal_pdf(path: Path, text: str) -> None:
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n"
        + stream
        + b"\nendstream",
    ]
    payload = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, body in enumerate(objects, start=1):
        offsets.append(len(payload))
        payload.extend(f"{number} 0 obj\n".encode("ascii"))
        payload.extend(body)
        payload.extend(b"\nendobj\n")
    xref_offset = len(payload)
    payload.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    payload.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        payload.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    payload.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(payload)
